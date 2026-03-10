"""生成91天复习课表 Word 文档（2000词课程示例）"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---- 页面设置 ----
section = doc.sections[0]
section.page_width = Inches(11)
section.page_height = Inches(8.5)
section.left_margin = Cm(2)
section.right_margin = Cm(2)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

# ---- 标题 ----
title = doc.add_heading('91天词汇复习全程计划', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

sub = doc.add_paragraph('以2000词课程（10课，每课200词，15天学习期）为例 | 实际词量会扣除已掌握单词')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(10)
sub.runs[0].font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

doc.add_paragraph()

# 阶段颜色
PHASE_COLORS = {
    'learn':  RGBColor(0xdb, 0xed, 0xff),  # 蓝
    'review': RGBColor(0xd1, 0xfa, 0xe5),  # 绿
    'rest':   RGBColor(0xf3, 0xf4, 0xf6),  # 灰
    'test':   RGBColor(0xff, 0xe4, 0xe6),  # 红
}
PHASE_LABELS = {
    'learn':  '📘 学习',
    'review': '🔄 复习',
    'rest':   '😴 休息',
    'test':   '📝 检测',
}

# ---- 数据 ----
schedule = []

# Phase 1: 学习期
LESSON_WORDS = 200
learn_data = [
    (1,  'learn', [1],       200,  'L1'),
    (2,  'learn', [1,2],     400,  'L2'),
    (3,  'learn', [2,3],     400,  'L3'),
    (4,  'learn', [1,3,4],   600,  'L4'),
    (5,  'learn', [2,4,5],   600,  'L5'),
    (6,  'rest',  [],          0,  '假期'),
    (7,  'rest',  [],          0,  '假期'),
    (8,  'rest',  [],          0,  '假期'),
    (9,  'rest',  [],          0,  '假期'),
    (10, 'rest',  [],          0,  '假期'),
    (11, 'learn', [3,5,6],   600,  'L6'),
    (12, 'learn', [4,6,7],   600,  'L7'),
    (13, 'learn', [5,7,8],   600,  'L8'),
    (14, 'learn', [6,8,9],   600,  'L9'),
    (15, 'learn', [7,9,10],  600,  'L10'),
]
# Phase 2: 抗遗忘
anti = [
    (16, 'review', [1,5,8,9],    800, '抗遗忘D1'),
    (17, 'review', [2,6,9,10],   800, '抗遗忘D2'),
    (18, 'review', [3,7,10],     600, '抗遗忘D3'),
    (19, 'review', [4,8],        400, '抗遗忘D4'),
]
# Phase 3: 第二轮全面复习
round2 = [(20+i, 'review', [i+1], 200, f'第二轮L{i+1}') for i in range(10)]

# Phase 4: 休息
rest1 = [(30+i, 'rest', [], 0, '休息') for i in range(5)]

# Phase 5: 第三轮全面复习
round3 = [(35+i, 'review', [i+1], 200, f'第三轮L{i+1}') for i in range(10)]

# Phase 6: 休息+检测1
rest2 = [(45+i, 'rest', [], 0, '休息') for i in range(13)]
test1 = [(58, 'test', list(range(1,11)), 2000, '综合检测1')]

# Phase 7: 休息+检测2
rest3 = [(59+i, 'rest', [], 0, '休息') for i in range(32)]
test2 = [(91, 'test', list(range(1,11)), 2000, '综合检测2')]

rows = learn_data + anti + round2 + rest1 + round3 + rest2 + test1 + rest3 + test2

# ---- 表格 ----
headers = ['DAY', '类型', '阶段说明', '复习课程', '复习词数（上限）']
col_widths = [Cm(1.5), Cm(1.8), Cm(3.2), Cm(8), Cm(2.8)]

table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# 表头
hrow = table.rows[0]
hrow.height = Cm(0.8)
for i, (h, w) in enumerate(zip(headers, col_widths)):
    cell = hrow.cells[i]
    cell.width = w
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(9)
    # 表头背景色
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1a56db')
    tcPr.append(shd)
    run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)

# 数据行
rest_group = 0
for day, dtype, lessons, words, label in rows:
    row = table.add_row()
    row.height = Cm(0.6)

    # 决定显示
    lesson_str = ', '.join(f'L{l}' for l in lessons) if lessons else '—'
    if dtype == 'rest':
        lesson_str = '—'
        words_str = '—'
    elif dtype == 'test':
        lesson_str = 'L1~L10（全部）'
        words_str = f'{words:,}词（全量）'
    else:
        words_str = f'~{words}词'

    values = [str(day), PHASE_LABELS[dtype], label, lesson_str, words_str]
    fill = PHASE_COLORS[dtype].rgb if hasattr(PHASE_COLORS[dtype], 'rgb') else ''.join(f'{c:02x}' for c in PHASE_COLORS[dtype])
    r, g, b = PHASE_COLORS[dtype][0], PHASE_COLORS[dtype][1], PHASE_COLORS[dtype][2]
    fill_hex = f'{r:02x}{g:02x}{b:02x}'

    for i, (val, w) in enumerate(zip(values, col_widths)):
        cell = row.cells[i]
        cell.width = w
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 3 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(val)
        run.font.size = Pt(8.5)
        if dtype == 'test':
            run.bold = True
        # 背景色
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_hex)
        tcPr.append(shd)

# ---- 图例 ----
doc.add_paragraph()
legend = doc.add_paragraph()
legend.alignment = WD_ALIGN_PARAGRAPH.LEFT
legend.add_run('图例：').bold = True
legend.add_run('  📘 学习期（蓝）  🔄 复习（绿）  😴 休息（灰）  📝 综合检测（红）')
legend.runs[-1].font.size = Pt(9)

# ---- 说明 ----
note = doc.add_paragraph()
note.add_run('注：').bold = True
note.add_run(
    '① 复习词数为理论上限，实际词数 = 理论词数 − 已掌握词数（连续答对3次自动扣除）。\n'
    '② 综合检测复习全部2000词，请做好充分准备。\n'
    '③ 本课表适用于「10天2000词」班型（含公共假期5天，学习期共15天）。'
)
note.runs[-1].font.size = Pt(9)
note.runs[-1].font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

out = r'e:\trae\1word\91天复习全程计划.docx'
doc.save(out)
print(f'已生成: {out}')
