# -*- coding: utf-8 -*-
"""
@file extract-quiz-questions.py
@desc 从Word文档提取词汇量测试题目并生成JSON文件
@see PRD: docs/PRD.md#VocabularyQuiz
"""

import json
import re
from docx import Document

def extract_quiz_questions(docx_path: str, output_path: str):
    """
    从Word文档中提取词汇量测试题目
    
    Args:
        docx_path: Word文档路径
        output_path: 输出JSON文件路径
    """
    doc = Document(docx_path)
    
    questions = []
    current_question = None
    question_no = 0
    
    # 解析段落获取题目
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 匹配题目格式: "1. deliberate [单选题] *"
        match = re.match(r'^(\d+)\.\s*(\w+)\s*\[单选题\]', text)
        if match:
            question_no = int(match.group(1))
            word = match.group(2)
            current_question = {
                'questionNo': question_no,
                'word': word,
                'optionA': '',
                'optionB': '',
                'optionC': '',
                'optionD': '以上都不对',
                'optionE': '不认识',
                'correctOption': '',
                'difficulty': 1  # 默认难度
            }
    
    # 解析表格获取选项
    # 第一个表格是年级选择，从第二个表格开始是题目选项
    for i, table in enumerate(doc.tables):
        if i == 0:
            continue  # 跳过年级选择表格
        
        question_index = i - 1  # 对应题目索引
        if question_index >= 50:
            break
        
        options = []
        correct_option = ''
        
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if not cell_text:
                    continue
                
                # 匹配选项格式: "A. 脆弱的" 或 "B. 深思熟虑的(正确答案)"
                option_match = re.match(r'^([A-E])[\.\s、]+(.+)$', cell_text)
                if option_match:
                    option_letter = option_match.group(1)
                    option_content = option_match.group(2).strip()
                    
                    # 检查是否是正确答案
                    if '(正确答案)' in option_content or '（正确答案）' in option_content:
                        correct_option = option_letter
                        option_content = option_content.replace('(正确答案)', '').replace('（正确答案）', '').strip()
                    
                    options.append({
                        'letter': option_letter,
                        'content': option_content
                    })
        
        # 构建题目对象
        if options:
            question = {
                'questionNo': question_index + 1,
                'word': '',
                'optionA': '',
                'optionB': '',
                'optionC': '',
                'optionD': '以上都不对',
                'optionE': '不认识',
                'correctOption': correct_option,
                'difficulty': 1
            }
            
            for opt in options:
                letter = opt['letter']
                content = opt['content']
                question[f'option{letter}'] = content
            
            questions.append(question)
    
    # 从段落中提取单词并匹配
    word_list = []
    for para in doc.paragraphs:
        text = para.text.strip()
        match = re.match(r'^(\d+)\.\s*(\w+)\s*\[单选题\]', text)
        if match:
            word_list.append({
                'no': int(match.group(1)),
                'word': match.group(2)
            })
    
    # 将单词匹配到题目
    for i, q in enumerate(questions):
        if i < len(word_list):
            q['word'] = word_list[i]['word']
            q['questionNo'] = word_list[i]['no']
    
    # 保存JSON
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(questions, f, ensure_ascii=False, indent=2)
    
    print(f'成功提取 {len(questions)} 道题目，已保存到 {output_path}')
    
    # 打印前5道题目预览
    print('\n题目预览:')
    for q in questions[:5]:
        print(f"{q['questionNo']}. {q['word']}")
        print(f"   A. {q['optionA']}")
        print(f"   B. {q['optionB']}")
        print(f"   C. {q['optionC']}")
        print(f"   D. {q['optionD']}")
        print(f"   E. {q['optionE']}")
        print(f"   正确答案: {q['correctOption']}")
        print()
    
    return questions

if __name__ == '__main__':
    import os
    
    # 获取脚本所在目录
    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.dirname(script_dir)
    
    docx_path = os.path.join(os.path.dirname(project_root), '词汇量测试（2000词）-有答案版本.docx')
    output_path = os.path.join(script_dir, 'quiz-questions.json')
    
    print(f'正在从 {docx_path} 提取题目...')
    extract_quiz_questions(docx_path, output_path)
